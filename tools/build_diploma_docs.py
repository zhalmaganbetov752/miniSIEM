# -*- coding: utf-8 -*-
"""Generate diploma PZ package for Жалмаганбетов Е.М."""
from __future__ import annotations

import re
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "ПЗ"
SRC_PZ = OUT_DIR / "ПЗ.docx"
if not SRC_PZ.exists():
    SRC_PZ = ROOT / "ПЗ.docx"
MD_SUPP = OUT_DIR / "ПЗ_дополнение_введение_главы3-5_заключение.md"
if not MD_SUPP.exists():
    MD_SUPP = ROOT / "ПЗ_дополнение_введение_главы3-5_заключение.md"

STUDENT = {
    "fio": "Жалмаганбетов Еламан Манарбекулы",
    "fio_short": "Жалмаганбетов Е.М.",
    "group": "СИБ 23-2с",
    "supervisor": "Винтерголлер И.Г.",
    "supervisor_full": "ст. преп. кафедры «Кибербезопасности и искусственного интеллекта» Винтерголлер И.Г.",
    "reviewer": "Хасенова Айгерим Асхатовна",
    "defense": "10.06.2026",
    "year": "2026",
    "city": "Караганда",
}

TOPIC_RU = (
    "Разработка программного средства анализа журналов безопасности "
    "операционной системы на рабочей станции"
)
TOPIC_KZ = (
    "Жұмыс станциясында операциялық жүйе қауіпсіздік журналдарын "
    "талдауға арналған бағдарламалық құралды әзірлеу"
)
SPECIALTY = "6B06301 — Системы информационной безопасности"
UNIVERSITY = "НАО «Карагандинский технический университет имени Абылкаса Сагинова»"
DEPARTMENT = "Кафедра «Кибербезопасности и искусственного интеллекта»"
MINISTRY = "Министерство науки и высшего образования Республики Казахстан"


def _style_body(p, *, bold: bool = False, center: bool = False, indent: bool = True) -> None:
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if indent and not center:
        pf.first_line_indent = Cm(1.25)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
        r.bold = bold or r.bold


def _add(doc: Document, text: str, *, bold: bool = False, center: bool = False, indent: bool = True) -> None:
    p = doc.add_paragraph(text)
    _style_body(p, bold=bold, center=center, indent=indent)


def _parse_md(path: Path) -> dict[str, list[str]]:
    text = path.read_text(encoding="utf-8")
    out: dict[str, list[str]] = {}
    key: str | None = None
    buf: list[str] = []
    for line in text.splitlines():
        if line.startswith("### "):
            if key:
                out[key] = [ln for ln in buf if ln.strip() and not ln.startswith("```")]
            key = line[4:].strip()
            buf = []
            continue
        if line.startswith("## "):
            if key:
                out[key] = [ln for ln in buf if ln.strip() and not ln.startswith("```")]
            key = line[3:].strip()
            buf = []
            continue
        if line.startswith("# ") or line.startswith("---"):
            continue
        if key is not None:
            buf.append(line)
    if key:
        out[key] = [ln for ln in buf if ln.strip() and not ln.startswith("```")]
    return out


def _p_text(p) -> str:
    return "".join(r.text for r in p.runs).strip()


def _extract_chapters_1_2(src: Path) -> list[tuple[str, bool]]:
    """Paragraphs from original PZ between ch1 and empty Заключение."""
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    with zipfile.ZipFile(src) as z:
        root = ET.fromstring(z.read("word/document.xml"))
    paras: list[str] = []
    for para in root.iter(W + "p"):
        t = "".join(x.text or "" for x in para.iter(W + "t")).strip()
        if t:
            paras.append(t)
    start = next(i for i, t in enumerate(paras) if t == "1. Предпроектные исследования")
    end = next(i for i, t in enumerate(paras) if t == "Заключение" and i > start)
    block = paras[start:end]
    result: list[tuple[str, bool]] = []
    for t in block:
        bold = bool(
            re.match(r"^\d+\.", t)
            or t.startswith("Таблица")
            or t in ("Критерий", "Аспект", "Параметр", "Компонент стека", "Технология", "Роль в системе", "Язык")
        )
        result.append((t, bold))
    return result


def _blocks_from_sections(sections: dict[str, list[str]]) -> list[tuple[str, bool]]:
    order = [
        "ВВЕДЕНИЕ",
        "3. ПРОЕКТИРОВАНИЕ ПРОГРАММНОГО КОМПЛЕКСА mini SIEM",
        "4. РЕАЛИЗАЦИЯ ПРОГРАММНОГО КОМПЛЕКСА",
        "5. ТЕСТИРОВАНИЕ И АПРОБАЦИЯ РЕЗУЛЬТАТОВ",
        "ЗАКЛЮЧЕНИЕ",
    ]
    blocks: list[tuple[str, bool]] = []
    for main in order:
        if main not in sections:
            continue
        blocks.append((main.replace("ВВЕДЕНИЕ", "Введение").replace("ЗАКЛЮЧЕНИЕ", "Заключение"), True))
        m = re.match(r"^(\d+)\.", main)
        if m:
            num = m.group(1)
            for sk in sorted(k for k in sections if re.match(rf"^{num}\.\d", k)):
                blocks.append((sk, True))
                for line in sections[sk]:
                    blocks.append((line, False))
        else:
            for line in sections[main]:
                blocks.append((line, False))
    return blocks


def _literature_from_src(src: Path) -> list[str]:
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    with zipfile.ZipFile(src) as z:
        root = ET.fromstring(z.read("word/document.xml"))
    paras: list[str] = []
    for para in root.iter(W + "p"):
        t = "".join(x.text or "" for x in para.iter(W + "t")).strip()
        if t:
            paras.append(t)
    # Use the last heading (skip TOC line like "Список ...22")
    indices = [
        i
        for i, t in enumerate(paras)
        if t == "Список использованной литературы"
        or (t.startswith("Список использованной литературы") and not any(c.isdigit() for c in t[len("Список использованной литературы") :]))
    ]
    if not indices:
        indices = [i for i, t in enumerate(paras) if "Список использованной литературы" in t]
    i = indices[-1]
    refs = [t for t in paras[i + 1 :] if t and re.match(r"^\d+\.", t) or t.startswith("ISO")]
    return [paras[i]] + refs


def build_title() -> Path:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    _add(doc, MINISTRY, center=True, indent=False)
    _add(doc, UNIVERSITY, center=True, indent=False)
    _add(doc, DEPARTMENT, center=True, indent=False)
    doc.add_paragraph()
    _add(doc, "ДИПЛОМНЫЙ ПРОЕКТ", bold=True, center=True, indent=False)
    _add(doc, f"на тему: «{TOPIC_RU}»", center=True, indent=False)
    doc.add_paragraph()
    doc.add_paragraph()
    for line in (
        f"Выполнил: студент группы {STUDENT['group']}",
        STUDENT["fio"],
        "",
        f"Руководитель: {STUDENT['supervisor_full']}",
        "",
        f"Рецензент: {STUDENT['reviewer']}",
    ):
        p = doc.add_paragraph(line or " ")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.first_line_indent = Cm(0)
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(14)
    doc.add_paragraph()
    _add(doc, f"{STUDENT['city']} {STUDENT['year']}", center=True, indent=False)

    path = OUT_DIR / "01_Титульный_лист.docx"
    doc.save(str(path))
    return path


def build_assignment() -> Path:
    doc = Document()
    _add(doc, MINISTRY, center=True, indent=False)
    _add(doc, UNIVERSITY, center=True, indent=False)
    _add(doc, "ЛИСТ ЗАДАНИЯ", bold=True, center=True, indent=False)
    _add(doc, "на дипломный проект", center=True, indent=False)
    doc.add_paragraph()

    lines = [
        f"Студент: {STUDENT['fio']}, группа {STUDENT['group']}",
        f"Специальность: {SPECIALTY}",
        f"Тема: «{TOPIC_RU}»",
        "",
        "Исходные данные: журналы Windows Event Log (Security, System, Application); "
        "опционально Microsoft-Windows-Sysmon/Operational; архивы EVTX; "
        "техническая документация Microsoft; материалы производственной практики.",
        "",
        "Содержание расчётно-пояснительной записки:",
        "— введение;",
        "— предпроектные исследования (предметная область, аналоги, постановка задачи);",
        "— выбор технологий и средств разработки;",
        "— проектирование программного комплекса mini SIEM;",
        "— реализация программного комплекса;",
        "— тестирование и апробация результатов;",
        "— заключение;",
        "— список использованной литературы;",
        "— приложения (листинг программы, нормоконтроль).",
        "",
        "Перечень графического материала: архитектурная схема; скриншоты GUI; "
        "схема БД; примеры отчётов; демонстрационные листы.",
        "",
        "Дата выдачи задания: «___» __________ 2026 г.",
        "Срок сдачи законченного проекта: «10» июня 2026 г.",
        "",
        f"Руководитель {STUDENT['supervisor']} _________________",
        f"Студент {STUDENT['fio_short']} _________________",
    ]
    for ln in lines:
        _add(doc, ln, indent=bool(ln and not ln.startswith("—")))

    path = OUT_DIR / "02_Лист_задания.docx"
    doc.save(str(path))
    return path


def build_annotation() -> Path:
    doc = Document()
    _add(doc, "АННОТАЦИЯ", bold=True, center=True, indent=False)
    _add(doc, f"Тақырыбы: «{TOPIC_KZ}»", indent=False)
    _add(doc, f"Тақырыбы (орысша): «{TOPIC_RU}»", indent=False)
    _add(
        doc,
        f"Орындаған: {STUDENT['group']} тобы, 6B06301 мамандығы, {STUDENT['fio']}. "
        f"Жетекшісі: {STUDENT['supervisor']}. "
        f"{UNIVERSITY}, «Кибербезопасность және жасанды интеллект» кафедрасы.",
        indent=False,
    )
    doc.add_paragraph()
    kz = (
        "Дипломдық жобаның мақсаты — Windows жұмыс станциясындағы қауіпсіздік "
        "оқиғаларын жинақтау, сақтау және талдауға арналған mini SIEM бағдарламалық "
        "кешенін әзірлеу. Жүйе журналдардан оқиғаларды жинайды, SQLite дерекқорына "
        "жазады, YAML ережелері бойынша детекторларды қолданады, дабылдар құрайды "
        "және HTML/JSON есептерін шығарады. Нәтиже — оқу-зерттеу және демонстрациялық "
        "пайдалануға арналған жұмыс істейтін прототип."
    )
    ru = (
        "Цель дипломного проекта — разработка программного комплекса mini SIEM для "
        "сбора, хранения и анализа событий информационной безопасности на рабочей "
        "станции Windows. Система опрашивает журналы событий, сохраняет данные в SQLite, "
        "применяет правила детектирования, формирует алерты и отчёты. "
        f"Специальность 6B06301 «Системы информационной безопасности». "
        f"Объём пояснительной записки — 45–60 страниц. Ключевые слова: SIEM, Windows Event Log, "
        "Sysmon, детектирование, корреляция, SQLite, Python."
    )
    _add(doc, kz)
    _add(doc, ru)
    path = OUT_DIR / "03_Аннотация.docx"
    doc.save(str(path))
    return path


def _split_blocks(all_blocks: list[tuple[str, bool]]) -> tuple[list, list, list]:
    intro: list[tuple[str, bool]] = []
    ch35: list[tuple[str, bool]] = []
    concl: list[tuple[str, bool]] = []
    mode: str | None = None
    for text, bold in all_blocks:
        if text == "Введение":
            mode = "intro"
        elif text.startswith("3. ПРОЕКТИРОВАНИЕ"):
            mode = "ch35"
        elif text == "Заключение":
            mode = "concl"
        if mode == "intro":
            intro.append((text, bold))
        elif mode == "ch35":
            ch35.append((text, bold))
        elif mode == "concl":
            concl.append((text, bold))
    return intro, ch35, concl


def build_full_pz() -> Path:
    sections = _parse_md(MD_SUPP)
    ch12 = _extract_chapters_1_2(SRC_PZ)
    all_blocks = _blocks_from_sections(sections)
    intro, ch35, concl = _split_blocks(all_blocks)
    literature = _literature_from_src(SRC_PZ)

    toc_lines = [
        "Введение",
        "1. Предпроектные исследования",
        "1.1. Анализ предметной области",
        "1.2. Анализ существующих разработок",
        "1.3. Постановка задачи",
        "2. Выбор технологий и средств разработки",
        "3. Проектирование программного комплекса mini SIEM",
        "3.1. Общая архитектура и конвейер обработки событий",
        "3.2. Проектирование хранилища данных",
        "3.3. Проектирование подсистемы сбора и нормализации",
        "3.4. Проектирование подсистемы правил и корреляции",
        "3.5. Проектирование пользовательского интерфейса и отчётности",
        "3.6. Проектирование графического материала дипломного проекта",
        "4. Реализация программного комплекса",
        "4.1. Модуль сбора событий Windows",
        "4.2. Модуль хранения и обслуживания базы данных",
        "4.3. Модуль правил детектирования",
        "4.4. Графическое приложение аналитика",
        "4.5. Импорт EVTX, отчёты и поставка",
        "5. Тестирование и апробация результатов",
        "5.1. Методика и стенд тестирования",
        "5.2. Функциональное тестирование",
        "5.3. Тестирование правил детектирования",
        "5.4. Оценка производительности и ограничений",
        "5.5. Соответствие критериям приёмки и готовность к защите 10.06.2026",
        "Заключение",
        "Список использованной литературы",
        "Приложение А. Листинг программы",
        "Приложение Б. Перечень замечаний нормоконтроллера",
    ]

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    _add(doc, "СОДЕРЖАНИЕ", bold=True, center=True, indent=False)
    for ln in toc_lines:
        _add(doc, ln, indent=False)
    doc.add_page_break()

    for block in (intro, ch12, ch35, concl):
        for text, bold in block:
            _add(doc, text, bold=bold)

    doc.add_page_break()
    _add(doc, "Список использованной литературы", bold=True, indent=False)
    for item in literature[1:]:
        _add(doc, item, indent=False)
    for item in [
        "21. NIST SP 800-92. Guide to Computer Security Log Management. URL: https://csrc.nist.gov/publications/detail/sp/800-92/final (дата обращения: 25.05.2026).",
        "22. Python Software Foundation. Python 3.11 Documentation. URL: https://docs.python.org/3.11/ (дата обращения: 25.05.2026).",
        "23. Qt Project. Qt for Python (PySide6). URL: https://doc.qt.io/qtforpython/ (дата обращения: 25.05.2026).",
    ]:
        _add(doc, item, indent=False)

    path = OUT_DIR / "04_ПЗ_Жалмаганбетов_Е.М._полная.docx"
    doc.save(str(path))
    return path


def build_screenshots_guide() -> Path:
    text = f"""# Скриншоты для диплома — {STUDENT['fio']}

Защита: {STUDENT['defense']}. Вставляйте рисунки в разделы 3–5 с подписями «Рисунок N — …».

## Обязательные (минимум 6)

| № | Что снять | Где в ПЗ | Как получить |
|---|-----------|----------|--------------|
| 1 | Главное окно, вкладка **Dashboard** (виджеты: события за 24 ч, алерты, топ EventID) | п. 3.5, 4.4 | `python -m minisiem` → открыть Dashboard после 5–10 мин работы |
| 2 | Вкладка **Events**: таблица событий + панель фильтров (канал Security, EventID 4625) | п. 4.4 | Применить фильтр, чтобы было видно 5+ строк |
| 3 | Вкладка **Alerts**: список алертов с severity (high/critical), статус **new** | п. 3.4, 5.3 | Сгенерировать неудачные входы или импорт EVTX |
| 4 | Карточка/окно **деталей алерта** + кнопка экспорта HTML-отчёта | п. 4.5 | Открыть любой алерт → Export / отчёт |
| 5 | Вкладка **Rules**: список правил, WIN-SEC-4625-BURST включено, параметры threshold/window | п. 3.4, 4.3 | Вкладка Rules |
| 6 | **HTML-отчёт** в браузере (сводка или по инциденту) | п. 4.5 | Экспорт → открыть `export/minisiem-report.html` |

## Желательные (+4)

| № | Что снять | Где в ПЗ |
|---|-----------|----------|
| 7 | **Импорт EVTX**: диалог выбора файла + прогресс | п. 4.5 |
| 8 | Вкладка **Settings**: интервал опроса, dry-run, путь к БД, retention | п. 4.4 |
| 9 | **Просмотр события** 4625: message + parsed IP/account | п. 4.1 |
| 10 | Вкладка **Actions** (журнал реакций, dry-run) | п. 3.2, 4.4 |

## Для графических листов (не скриншот, а схема)

| № | Материал | Инструмент |
|---|----------|------------|
| Г1 | Архитектура mini SIEM (блоки: collector → db → rules → GUI) | draw.io / Visio |
| Г2 | ER-диаграмма таблиц events, alerts, actions | draw.io |
| Г3 | Диаграмма последовательности: сбор → детект → алерт | draw.io |

## Как получить алерты для скриншотов

**Вариант А (тестовая ВМ):** 5+ неудачных входов RDP/SMB на учётную запись (только своя ВМ).

**Вариант Б:** Импорт готового EVTX с событиями 4625/4624 (Events → Import EVTX).

**Вариант В:** PowerShell с `-EncodedCommand` при установленном Sysmon → алерт WIN-SYSMON-PS-ENC.

## Требования к оформлению

- Разрешение не ниже 1920×1080, масштаб Windows 100%.
- Обрезать лишнее (только окно приложения).
- Подпись: «Рисунок 3 — Вкладка Alerts программного комплекса mini SIEM».
- В тексте ссылка: «… представлено на рисунке 3».
"""
    path = OUT_DIR / "05_Скриншоты_инструкция.md"
    path.write_text(text, encoding="utf-8")
    return path


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    paths = [
        build_title(),
        build_assignment(),
        build_annotation(),
        build_full_pz(),
        build_screenshots_guide(),
    ]
    print("Created:")
    for p in paths:
        print(" ", p)


if __name__ == "__main__":
    main()
