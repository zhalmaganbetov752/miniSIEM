# -*- coding: utf-8 -*-
"""Remove duplicate and redundant paragraphs from ПЗ.docx."""
from __future__ import annotations

import re
import shutil
from collections import defaultdict
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "ПЗ" / "ПЗ.docx"
OUT = ROOT / "ПЗ" / "ПЗ.docx"
BACKUP = ROOT / "ПЗ" / "ПЗ_before_dedupe.docx"

# Явно ошибочные / дублирующие абзацы (индексы до удаления; пересчитываются при batch-delete)
EXPLICIT_SNIPPETS = [
    # Дублирующий блок «реализация всех модулей» в 4.4
    "Реализация начинается с точки входа __main__.py",
    "В collector_win.py функция poll_channel выбирает ветку чтения",
    "Модуль db.py (~930 строк) концентрирует работу с SQLite. insert_event записывает",
    "Модуль rules.py загружает YAML, строит список RuleDef",
    "Графическое приложение app.py (~2900 строк) использует QMainWindow",
    # Чужие для 4.5 абзацы про главу 5
    "Раздел 5.5 готовности к защите перечисляет комплект документов",
    "Раздел 5.4 оценки производительности фиксирует качественные наблюдения",
    "Раздел 5.3 дополняется описанием negative testing",
    "Глава 5 дополняется таблицей соответствия сценариев атак",
    # Дублирующие вставки расширения в 3.1 (после оригинального текста)
    "Архитектура mini SIEM построена по конвейерному принципу: источник (Windows",
    "Класс MiniSIEMApp в app.py выступает оркестратором",
    "Диаграмма архитектуры (графический материал дипломного проекта) отражает",
    "Конвейер включает семь этапов: Collection, Normalization, Storage",
    "Расширяемость архитектуры: новый детектор — функция в rules.py",
    "При эксплуатации все компоненты размещаются на одной рабочей станции; сетевое",
    "Сбор (Collection) выполняется таймером QTimer",
    "Логическая схема на рисунке 1 показывает поток: Event Log/EVTX",
    "Модульность позволяет заменить SQLite на PostgreSQL",
    "Response-этап конвейера реализован таблицей actions",
    # Повтор телеком-сети (сокращённый вариант)
    "Телекоммуникационная сеть оператора в логической модели представляет собой иерархию уровней: магистраль, агрегация, доступ, сервисные платформы (аутентификация AAA",
    # Дубли 4.5 после таблицы 8 (повтор 504-506)
    "Импорт EVTX в evtx_import.py читает XML-записи из архива, ограничивает",
    "Модуль report.py формирует HTML-отчёт build_html_report",
    "Сборка PyInstaller по packaging/minisiem.spec включает minisiem",
    # Повторы в 4.1 / 4.4
    "Точка входа __main__.py вызывает main() класса MiniSIEMApp",
    "Функция poll_channel в collector_win.py принимает имя канала",
    "Класс MiniSIEMApp (~2900 строк в app.py) создаёт QMainWindow",
    "Нормализация EventID выполняется операцией побитового И с маской 0xFFFF",
    # Почти дубль в 1.2
    "Аналитика ложных срабатываний в зрелых SIEM приводит к дисциплинам тюнинга",
    # Повторы в 1.2 / 1.3
    "В результате анализа существующих разработок сделан вывод:",
    "Ожидаемый результат: работоспособный mini SIEM, подтверждающий на защите достижение",
    # Лишние meta-абзацы в конце глав
    "Глава 4 дополняется перечислением ключевых классов app.py",
    "Реализация rules.py включает вспомогательные функции _parse_ts",
    "Тестирование включает проверку отображения часового пояса: событие с ts_utc",
]

SIMILARITY = 0.72
PREFIX_LEN = 55
MIN_SIM_LEN = 50


def _p_text(p) -> str:
    return "".join(r.text for r in p.runs).strip()


def _norm(s: str) -> str:
    s = s.lower().replace("ё", "е")
    return re.sub(r"\s+", " ", s).strip()


def _delete_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def _find_remove_indices(doc: Document) -> list[int]:
    remove: set[int] = set()
    items = [(i, _p_text(p)) for i, p in enumerate(doc.paragraphs)]

    # 1) Точные дубликаты (символы диаграммы и пр.)
    seen_exact: dict[str, int] = {}
    for i, t in items:
        if not t:
            continue
        if t in seen_exact:
            remove.add(i)
        else:
            seen_exact[t] = i

    # 2) Явные фрагменты
    for i, t in items:
        if not t or i in remove:
            continue
        for snippet in EXPLICIT_SNIPPETS:
            if t.startswith(snippet) or snippet in t:
                remove.add(i)
                break

    # 3) Похожие абзацы: одинаковое начало + ratio >= SIMILARITY (оставить первый)
    buckets: dict[str, list[tuple[int, str]]] = defaultdict(list)
    for i, t in items:
        if not t or len(t) < MIN_SIM_LEN or i in remove:
            continue
        buckets[_norm(t)[:PREFIX_LEN]].append((i, t))

    for group in buckets.values():
        if len(group) < 2:
            continue
        group.sort(key=lambda x: x[0])
        kept: list[tuple[int, str, str]] = []
        for i, t in group:
            n = _norm(t)
            is_dup = False
            for _, _, kn in kept:
                if SequenceMatcher(None, kn, n).ratio() >= SIMILARITY:
                    remove.add(i)
                    is_dup = True
                    break
            if not is_dup:
                kept.append((i, t, n))

    # 4) Повтор «Связь предметной области» — короткий вариант
    for i, t in items:
        if i in remove or not t:
            continue
        if t.startswith("Связь предметной области с дипломом выражается") and len(t) < 400:
            remove.add(i)

    # 5) Дубли «Объект автоматизации» / «Предмет разработки» если есть полный блок ТЗ
    for i, t in items:
        if i in remove or not t:
            continue
        if t.startswith("Объект автоматизации:") and len(t) < 150:
            remove.add(i)

    return sorted(remove, reverse=True)


def main() -> None:
    shutil.copy2(SRC, BACKUP)
    doc = Document(str(SRC))
    before = sum(1 for p in doc.paragraphs if _p_text(p))
    before_chars = sum(len(_p_text(p)) for p in doc.paragraphs if _p_text(p))

    remove = _find_remove_indices(doc)
    print(f"Removing {len(remove)} paragraphs (indices before delete): {sorted(remove, reverse=True)[:40]}...")

    for idx in remove:
        _delete_paragraph(doc.paragraphs[idx])

    after = sum(1 for p in doc.paragraphs if _p_text(p))
    after_chars = sum(len(_p_text(p)) for p in doc.paragraphs if _p_text(p))

    out = OUT
    try:
        doc.save(str(out))
    except PermissionError:
        out = ROOT / "ПЗ" / "ПЗ_без_дублей.docx"
        doc.save(str(out))

    # sync copy
    copy = ROOT / "ПЗ" / "ПЗ_расширено_50стр.docx"
    if out == OUT:
        shutil.copy2(out, copy)

    print(f"Paragraphs: {before} -> {after} (-{before - after})")
    print(f"Chars: {before_chars} -> {after_chars} (-{before_chars - after_chars})")
    print(f"Backup: {BACKUP}")
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
