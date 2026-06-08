# -*- coding: utf-8 -*-
"""Расширение объёма пояснительной записки ПЗ.docx."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "tools"))
from pz_expansion_text import EXPANSIONS  # noqa: E402
from pz_expansion_extra import EXPANSIONS_EXTRA  # noqa: E402
from pz_expansion_extra2 import EXPANSIONS_EXTRA2  # noqa: E402

for _k, _v in EXPANSIONS_EXTRA.items():
    EXPANSIONS.setdefault(_k, []).extend(_v)
for _k, _v in EXPANSIONS_EXTRA2.items():
    EXPANSIONS.setdefault(_k, []).extend(_v)

SRC = ROOT / "ПЗ" / "ПЗ.docx"
OUT = ROOT / "ПЗ" / "ПЗ.docx"

# Порядок якорей: ключ → заголовок следующего раздела (startswith)
ANCHORS: list[tuple[str, str]] = [
    ("Введение", "1 Предпроектные"),
    ("1.1", "1.2"),
    ("1.2", "1.3"),
    ("1.3", "2 Выбор"),
    ("2.1", "2.2"),
    ("2.2", "2.3"),
    ("2.3", "2.4"),
    ("2.4", "2.5"),
    ("2.5", "3 Проектирование"),
    ("3.1", "3.2"),
    ("3.2", "3.3"),
    ("3.3", "3.4"),
    ("3.4", "3.5"),
    ("3.5", "3.6"),
    ("3.6", "4 Реализация"),
    ("4.1", "4.2"),
    ("4.2", "4.3"),
    ("4.3", "4.4"),
    ("4.4", "4.5"),
    ("4.5", "5 Тестирование"),
    ("5.1", "5.2"),
    ("5.2", "5.3"),
    ("5.3", "5.4"),
    ("5.4", "5.5"),
    ("5.5", "Заключение"),
    ("Заключение", "Список использованной"),
]


def _p_text(p) -> str:
    return "".join(r.text for r in p.runs).strip()


def _style_para(p) -> None:
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)


def _insert_before(doc: Document, anchor_paragraph, text: str) -> None:
    anchor = anchor_paragraph._element
    parent = anchor.getparent()
    new_p = OxmlElement("w:p")
    parent.insert(parent.index(anchor), new_p)
    para = Paragraph(new_p, doc._body)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    _style_para(para)


def _find_anchor(doc: Document, prefix: str) -> Paragraph | None:
    for p in doc.paragraphs:
        t = _p_text(p)
        if t.startswith(prefix):
            return p
    return None


def main() -> None:
    doc = Document(str(SRC))
    inserted = 0

    for key, next_prefix in ANCHORS:
        paragraphs = EXPANSIONS.get(key)
        if not paragraphs:
            continue
        anchor = _find_anchor(doc, next_prefix)
        if anchor is None:
            print(f"WARN: anchor not found for {key!r} -> {next_prefix!r}")
            continue
        for text in paragraphs:
            _insert_before(doc, anchor, text)
            inserted += 1

    out_path = OUT
    try:
        doc.save(str(out_path))
    except PermissionError:
        out_path = ROOT / "ПЗ" / "ПЗ_расширено.docx"
        doc.save(str(out_path))

    chars = sum(len(_p_text(p)) for p in doc.paragraphs if _p_text(p))
    print(f"Inserted {inserted} paragraphs")
    print(f"Total chars: {chars} (~{chars // 1800} pages est.)")
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
