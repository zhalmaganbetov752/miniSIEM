# -*- coding: utf-8 -*-
"""Умеренное расширение ПЗ: между старой и полной версией (~85–90 тыс. символов)."""
from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "tools"))
from pz_expansion_text import EXPANSIONS  # noqa: E402
from pz_expansion_extra import EXPANSIONS_EXTRA  # noqa: E402

SRC = ROOT / "ПЗ" / "ПЗ.docx"
OUT = ROOT / "ПЗ" / "ПЗ.docx"
BACKUP = ROOT / "ПЗ" / "ПЗ_старая_до_умеренного.docx"
TARGET_CHARS = 90_000
MAX_FROM_BASE = 3
MAX_FROM_EXTRA = 1

SKIP_PREFIXES = (
    "Глава ",
    "Раздел ",
    "дополняется",
    "Реализация начинается с точки входа",
)

ANCHORS: list[tuple[str, str]] = [
    ("Введение", "1 Предпроектные"),
    ("1.1", "1.2"),
    ("1.2", "1.3"),
    ("1.3", "2 Выбор"),
    ("2.1", "2.2"),
    ("2.2", "2.3"),
    ("2.3", "2.4"),
    ("2.4", "2.5"),
    ("2.5", "3 Проектирование"),
    ("3.1", "3.2"),
    ("3.2", "3.3"),
    ("3.3", "3.4"),
    ("3.4", "3.5"),
    ("3.5", "3.6"),
    ("3.6", "4 Реализация"),
    ("4.1", "4.2"),
    ("4.2", "4.3"),
    ("4.3", "4.4"),
    ("4.4", "4.5"),
    ("4.5", "5 Тестирование"),
    ("5.1", "5.2"),
    ("5.2", "5.3"),
    ("5.3", "5.4"),
    ("5.4", "5.5"),
    ("5.5", "Заключение"),
    ("Заключение", "Список использованной"),
]


def _p_text(p) -> str:
    return "".join(r.text for r in p.runs).strip()


def _norm(s: str) -> str:
    s = s.lower().replace("ё", "е")
    return re.sub(r"\s+", " ", s).strip()


def _style_para(p) -> None:
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)


def _insert_before(doc: Document, anchor_paragraph, text: str) -> None:
    anchor = anchor_paragraph._element
    parent = anchor.getparent()
    new_p = OxmlElement("w:p")
    parent.insert(parent.index(anchor), new_p)
    para = Paragraph(new_p, doc._body)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    _style_para(para)


def _find_anchor(doc: Document, prefix: str):
    for p in doc.paragraphs:
        if _p_text(p).startswith(prefix):
            return p
    return None


def _select_moderate_blocks(existing: set[str]) -> dict[str, list[str]]:
    """До 3 абзацев из EXPANSIONS + 1 из EXTRA, без дублей с текущим текстом."""
    out: dict[str, list[str]] = {}
    all_keys = set(EXPANSIONS) | set(EXPANSIONS_EXTRA)
    for key in all_keys:
        picked: list[str] = []
        seen: set[str] = set()

        def _try_add(text: str) -> bool:
            if any(text.startswith(p) for p in SKIP_PREFIXES):
                return False
            n = _norm(text)
            if n in existing or n in seen or len(text) < 80:
                return False
            picked.append(text)
            seen.add(n)
            return True

        base_count = 0
        for text in EXPANSIONS.get(key, []):
            if _try_add(text):
                base_count += 1
            if base_count >= MAX_FROM_BASE:
                break

        extra_count = 0
        for text in EXPANSIONS_EXTRA.get(key, []):
            if _try_add(text):
                extra_count += 1
            if extra_count >= MAX_FROM_EXTRA:
                break

        if picked:
            out[key] = picked
    return out


def main() -> None:
    if BACKUP.exists():
        shutil.copy2(BACKUP, SRC)
    else:
        shutil.copy2(SRC, BACKUP)
    doc = Document(str(SRC))
    existing = {_norm(_p_text(p)) for p in doc.paragraphs if _p_text(p)}
    before = sum(len(_p_text(p)) for p in doc.paragraphs if _p_text(p))

    blocks = _select_moderate_blocks(existing)
    inserted = 0

    for key, next_prefix in ANCHORS:
        if before >= TARGET_CHARS:
            break
        paragraphs = blocks.get(key, [])
        if not paragraphs:
            continue
        anchor = _find_anchor(doc, next_prefix)
        if anchor is None:
            print(f"WARN: anchor not found {key!r} -> {next_prefix!r}")
            continue
        for text in paragraphs:
            if before >= TARGET_CHARS:
                break
            if _norm(text) in existing:
                continue
            _insert_before(doc, anchor, text)
            existing.add(_norm(text))
            before += len(text)
            inserted += 1

    out = OUT
    try:
        doc.save(str(out))
    except PermissionError:
        out = ROOT / "ПЗ" / "ПЗ_умеренно.docx"
        doc.save(str(out))

    # sync readable copy name
    moderate_copy = ROOT / "ПЗ" / "ПЗ_умеренно.docx"
    if out == OUT:
        shutil.copy2(out, moderate_copy)

    after = sum(len(_p_text(p)) for p in doc.paragraphs if _p_text(p))
    print(f"Inserted {inserted} paragraphs")
    print(f"Chars: {after} (target ~{TARGET_CHARS}, old backup {BACKUP.name})")
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
