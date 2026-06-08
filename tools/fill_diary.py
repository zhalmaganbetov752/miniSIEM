"""Fill internship diary docx (KTU template). Run: python tools/fill_diary.py

Заполняется ~половина строк плана (14 из 29), остальные строки шаблона очищаются.
Геометрия таблицы (число строк/столбцов) не меняется; принудительное сжатие кегля отключено.
"""

from __future__ import annotations

from datetime import date, timedelta

from docx import Document
from docx.enum.text import WD_UNDERLINE


DOC_PATH = r"c:\Users\Admin\Documents\Практика\дневник.docx"

PRACTICE_START = date(2026, 4, 13)
PRACTICE_END = date(2026, 5, 16)

# Строк данных в шаблоне (без шапки): 29. Нужно ~в 2 раза меньше записей.
DATA_ROWS_TEMPLATE = 29
DATA_ROWS_FILL = 14


def _set_paragraph_text_with_underline(paragraph, full_text: str, underline_substr: str) -> None:
    paragraph.clear()
    idx = full_text.find(underline_substr)
    if idx == -1:
        paragraph.add_run(full_text)
        return
    if idx > 0:
        paragraph.add_run(full_text[:idx])
    r = paragraph.add_run(underline_substr)
    r.underline = WD_UNDERLINE.SINGLE
    tail = full_text[idx + len(underline_substr) :]
    if tail:
        paragraph.add_run(tail)


def _date_for_row(row_index: int, *, start: date, end: date, total_rows: int) -> date:
    span = (end - start).days
    if total_rows <= 1:
        return start
    off = (row_index * span) // (total_rows - 1)
    return start + timedelta(days=off)


def _reset_table_paragraph_format(table) -> None:
    """Сбрасывает явный кегль/интервал в ячейках (от прежних запусков со сжатием)."""
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                pf = p.paragraph_format
                pf.line_spacing = None
                pf.line_spacing_rule = None
                pf.space_before = None
                pf.space_after = None
                for r in p.runs:
                    r.font.size = None


def _build_rows() -> list[tuple[str, str, str, str, str]]:
    """14 строк: объединённые по смыслу блоки практики."""
    specs: list[tuple[str, str]] = [
        ("ОТ, ТБ, мануал; правила в офисе, эвакуация, электробезопасность.", "Инструктаж, мануал."),
        ("Конфиденциальность, ПДн; мониторинг сети и сервисов.", "Политики, наставник."),
        ("Тикеты: статусы; классификация обращений.", "Практика, кейсы."),
        ("Удалённая диагностика; CPE, Wi‑Fi (базовые шаги).", "Системы, консультация."),
        ("Скорость канала; признаки массового инцидента.", "Регламент, разбор."),
        ("Трассировка, DNS; подготовка к выезду.", "Утилиты, чек-лист."),
        ("Выезд к абоненту; выезд к корпоративному клиенту.", "Выезд."),
        ("Видеонаблюдение: монтаж; настройка регистратора.", "Объект, настройка."),
        ("Видео: пароли, сегментация; мониторинг, смена.", "Совет, смена."),
        ("Диагностика, эскалация; разбор сложного кейса.", "Анализ, разбор."),
        ("Заметки для диплома (mini SIEM); события для ИБ.", "Конспект, аналитика."),
        ("Эскалация, связь служб; нестабильная сессия.", "Наблюдение, тикет."),
        ("VPN: типовые отказы; черновик отчёта по практике.", "Регламент, самостоятельно."),
        ("Сводка компетенций; завершение практики с наставником.", "Самооценка, итог."),
    ]
    assert len(specs) == DATA_ROWS_FILL

    rows: list[tuple[str, str, str, str, str]] = []
    for i, (content, methods) in enumerate(specs):
        d = _date_for_row(i, start=PRACTICE_START, end=PRACTICE_END, total_rows=DATA_ROWS_FILL)
        rows.append((str(i + 1), content, methods, d.strftime("%d.%m.%Y"), "Выполнено"))
    return rows


def main() -> None:
    doc = Document(DOC_PATH)

    doc.paragraphs[38].text = (
        "в компании «ТрансТелеКом» (филиал в г. Астана, Республика Казахстан), "
        "отдел информационных технологий (ИТ)"
    )
    doc.paragraphs[37].text = " (нужное подчеркнуть) "

    base = "о прохождении практики: учебной, производственной, преддипломной, педагогической, исследовательской"
    _set_paragraph_text_with_underline(doc.paragraphs[36], base, "производственной")

    doc.paragraphs[43].text = "Начало практики «13» апреля 2026 г."
    doc.paragraphs[46].text = "Окончание практики «16» мая 2026 г."

    table = doc.tables[0]
    _reset_table_paragraph_format(table)

    entries = _build_rows()
    assert len(entries) == DATA_ROWS_FILL

    for i in range(DATA_ROWS_TEMPLATE):
        row = table.rows[i + 1]
        if i < len(entries):
            no, content, methods, when, mark = entries[i]
            row.cells[0].text = no
            row.cells[1].text = content
            row.cells[2].text = methods
            row.cells[3].text = when
            row.cells[4].text = mark
        else:
            for c in row.cells:
                c.text = ""

    doc.save(DOC_PATH)
    print("OK:", DOC_PATH, f"заполнено строк: {DATA_ROWS_FILL} из {DATA_ROWS_TEMPLATE}")


if __name__ == "__main__":
    main()
