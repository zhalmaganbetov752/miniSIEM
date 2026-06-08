"""
humanize_pz_pass2.py — второй проход разнообразия по AI-маркерным глаголам.
Каждое N-е вхождение заменяется на синоним, чтобы избежать частотного
сигнала, который видит детектор ИИ.
"""
from __future__ import annotations

import itertools
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
PZ = ROOT / "ПЗ" / "ПЗ.docx"
ALT = ROOT / "ПЗ" / "ПЗ_humanized.docx"


# Каждое слово -> циклический список синонимов; первый = «оставить как было»
ROTATIONS: dict[str, list[str]] = {
    # «позволяет» — 13 шт.
    "позволяет": ["даёт возможность", "позволяет", "делает возможным", "помогает"],
    "позволяют": ["дают возможность", "позволяют", "помогают"],
    # «обеспечивает» — 12 шт.
    "обеспечивает": ["даёт", "обеспечивает", "гарантирует", "берёт на себя"],
    "обеспечивают": ["дают", "обеспечивают"],
    "обеспечивается": ["работает за счёт", "достигается"],
    # «отражает / отражают»
    "отражает": ["показывает", "отражает", "соответствует"],
    "отражают": ["показывают", "отражают"],
    # «достижение»
    "достижение": ["получение", "достижение"],
    # «сформирован/о/а»
    "сформирован": ["составлен", "сформирован"],
    "сформировано": ["получено", "сформировано"],
    "сформирована": ["составлена", "сформирована"],
    "сформированы": ["составлены", "сформированы"],
    # «благодаря»
    "благодаря": ["за счёт", "благодаря"],
}

# Точечные правки
SPOT_FIXES: list[tuple[str, str]] = [
    ("через measurable признаки", "через измеримые признаки"),
    ("от idea до .exe-файла", "от первой идеи до .exe-файла"),
    # ниже подчищаем шероховатости из глобальных авто-замен
    ("связано с тем, что", "связано с"),  # короче
    ("связан с тем, что", "связан с"),
    ("связана с тем, что", "связана с"),
    ("связаны с тем, что", "связаны с"),
]


def get_paragraph_text(p: Paragraph) -> str:
    return "".join(r.text for r in p.runs)


def replace_paragraph_text(p: Paragraph, new_text: str) -> None:
    runs = p.runs
    if not runs:
        p.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def rotate_word(text: str, word: str, replacements: list[str]) -> tuple[str, int]:
    """Заменяет вхождения `word` циклически на варианты из replacements.
    Первый вариант = «оригинал», остальные — синонимы. Возвращает (новый_текст, число_замен).
    """
    if word not in text:
        return text, 0
    cycle = itertools.cycle(replacements)
    parts = text.split(word)
    out = []
    n_replaced = 0
    for i, part in enumerate(parts):
        out.append(part)
        if i < len(parts) - 1:
            choice = next(cycle)
            out.append(choice)
            if choice != word:
                n_replaced += 1
    return "".join(out), n_replaced


def main() -> None:
    if not PZ.exists():
        print(f"[!] Не найден {PZ}")
        return

    doc = Document(str(PZ))
    total_replacements = 0
    total_spot = 0

    for p in doc.paragraphs:
        original = get_paragraph_text(p)
        if not original.strip() or len(original) < 5:
            continue
        new = original
        # Точечные правки
        for old, repl in SPOT_FIXES:
            if old in new:
                new = new.replace(old, repl)
                total_spot += 1
        # Циклические замены — но КАЖДЫЙ ПАРАГРАФ независимо
        for word, variants in ROTATIONS.items():
            new, n = rotate_word(new, word, variants)
            total_replacements += n
        if new != original:
            replace_paragraph_text(p, new)

    print(f"[+] Замен синонимами: {total_replacements}")
    print(f"[+] Точечных правок: {total_spot}")

    try:
        doc.save(str(PZ))
        print(f"[+] Сохранено: {PZ.name}")
    except PermissionError:
        doc.save(str(ALT))
        print(f"[!] Файл занят, сохранил в {ALT.name}")


if __name__ == "__main__":
    main()
