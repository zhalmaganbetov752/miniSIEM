"""
humanize_pz_pass3.py — настоящая ротация синонимов поверх всего документа.
Используется ОБЩЕЕ состояние счётчиков, чтобы каждое вхождение слова
давало разный синоним, а не один и тот же при первом срабатывании в параграфе.
"""
from __future__ import annotations

import itertools
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
PZ = ROOT / "ПЗ" / "ПЗ.docx"
ALT = ROOT / "ПЗ" / "ПЗ_humanized.docx"


# Глобальные циклы — состояние сохраняется между параграфами
GLOBAL_CYCLES: dict[str, itertools.cycle] = {
    "даёт возможность": itertools.cycle([
        "позволяет",
        "даёт возможность",
        "делает возможным",
        "помогает",
        "позволяет",
        "это удобно для того, чтобы",
    ]),
    "даёт": itertools.cycle([
        "обеспечивает",
        "даёт",
        "берёт на себя",
        "гарантирует",
        "выполняет",
        "помогает",
    ]),
    "дают": itertools.cycle([
        "обеспечивают",
        "дают",
        "выполняют",
        "помогают",
    ]),
    "показывает": itertools.cycle([
        "показывает",
        "видно из",
        "иллюстрирует",
        "наглядно отражает",
    ]),
    "соответствует": itertools.cycle([
        "соответствует",
        "согласуется с",
        "отвечает требованиям",
        "выполняет требования",
    ]),
}


def get_paragraph_text(p: Paragraph) -> str:
    return "".join(r.text for r in p.runs)


def replace_paragraph_text(p: Paragraph, new_text: str) -> None:
    runs = p.runs
    if not runs:
        p.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def main() -> None:
    if not PZ.exists():
        print(f"[!] Не найден {PZ}")
        return

    doc = Document(str(PZ))
    counts: dict[str, int] = {k: 0 for k in GLOBAL_CYCLES}

    for p in doc.paragraphs:
        text = get_paragraph_text(p)
        if not text.strip():
            continue
        new = text
        # Разбираем по самым длинным ключам сначала, чтобы «даёт возможность»
        # не съедалось «даёт».
        for key in sorted(GLOBAL_CYCLES.keys(), key=len, reverse=True):
            if key not in new:
                continue
            parts = new.split(key)
            out = [parts[0]]
            cycle = GLOBAL_CYCLES[key]
            for piece in parts[1:]:
                replacement = next(cycle)
                if replacement != key:
                    counts[key] += 1
                out.append(replacement)
                out.append(piece)
            new = "".join(out)
        if new != text:
            replace_paragraph_text(p, new)

    print("[+] Циклические замены (глобальное состояние):")
    for k, v in counts.items():
        print(f"    {k!r}: {v} вхождений заменено синонимами")

    try:
        doc.save(str(PZ))
        print(f"[+] Сохранено: {PZ.name}")
    except PermissionError:
        doc.save(str(ALT))
        print(f"[!] Файл занят, сохранил в {ALT.name}")


if __name__ == "__main__":
    main()
