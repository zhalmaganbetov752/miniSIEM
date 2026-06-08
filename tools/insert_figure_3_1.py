# -*- coding: utf-8 -*-
"""Вставка Рисунок 3.1 в ПЗ.docx вместо placeholder."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "ПЗ" / "ПЗ.docx"
IMG = ROOT / "ПЗ" / "figures" / "Рисунок_3_1_Архитектура_mini_SIEM.png"
CAPTION = "Рисунок 3.1 – Архитектура mini SIEM (логические компоненты)"


def main() -> None:
    if not IMG.exists():
        raise SystemExit(f"Image not found: {IMG}")
    doc = Document(str(DOC))
    replaced = False
    for i, p in enumerate(doc.paragraphs):
        t = "".join(r.text for r in p.runs).strip()
        if t != "[Место для рисунка]":
            continue
        nxt = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else None
        if nxt and CAPTION in "".join(r.text for r in nxt.runs):
            p.clear()
            run = p.add_run()
            run.add_picture(str(IMG), width=Cm(15))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            replaced = True
            break
    if not replaced:
        raise SystemExit("Placeholder for Рисунок 3.1 not found")
    try:
        doc.save(str(DOC))
        print(f"Inserted into {DOC}")
    except PermissionError:
        alt = DOC.with_name("ПЗ_с_рисунком_3_1.docx")
        doc.save(str(alt))
        print(f"WARN: {DOC.name} открыт — сохранено в {alt}")


if __name__ == "__main__":
    main()
