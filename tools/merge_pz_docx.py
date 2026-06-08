"""Build ПЗ_полная.docx: chapters 1-2 from ПЗ.docx + new intro, ch3-5, conclusion."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "ПЗ.docx"
MD = ROOT / "ПЗ_дополнение_введение_главы3-5_заключение.md"
OUT = ROOT / "ПЗ_полная.docx"


def _paragraph_text(p) -> str:
    return "".join(r.text for r in p.runs).strip()


def _style_para(para: Paragraph, *, bold: bool = False) -> None:
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Pt(35)
    for run in para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.bold = bold or run.bold


def _insert_before(doc: Document, anchor_index: int, text: str, *, bold: bool = False) -> None:
    anchor = doc.paragraphs[anchor_index]._p
    parent = anchor.getparent()
    new_p = OxmlElement("w:p")
    parent.insert(parent.index(anchor), new_p)
    para = Paragraph(new_p, doc._body)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = bold
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Pt(35)


def _find(doc: Document, text: str, start: int = 0) -> int:
    for i, p in enumerate(doc.paragraphs):
        if i >= start and _paragraph_text(p) == text:
            return i
    return -1


def _parse_md(path: Path) -> dict[str, list[str]]:
    text = path.read_text(encoding="utf-8")
    out: dict[str, list[str]] = {}
    key: str | None = None
    buf: list[str] = []
    for line in text.splitlines():
        if line.startswith("## "):
            if key:
                out[key] = [ln for ln in buf if ln.strip() and not ln.startswith("```")]
            key = line[3:].strip()
            buf = []
            continue
        if line.startswith("# ") or line.startswith("---"):
            continue
        if key is not None:
            buf.append(line)
    if key:
        out[key] = [ln for ln in buf if ln.strip() and not ln.startswith("```")]
    return out


def _ordered_blocks(sections: dict[str, list[str]]) -> list[tuple[str, bool]]:
    """Return (text, is_heading) in document order."""
    order_main = [
        "ВВЕДЕНИЕ",
        "3. ПРОЕКТИРОВАНИЕ ПРОГРАММНОГО КОМПЛЕКСА mini SIEM",
        "4. РЕАЛИЗАЦИЯ ПРОГРАММНОГО КОМПЛЕКСА",
        "5. ТЕСТИРОВАНИЕ И АПРОБАЦИЯ РЕЗУЛЬТАТОВ",
        "ЗАКЛЮЧЕНИЕ",
    ]
    sub_order: list[str] = []
    for n in (3, 4, 5):
        subs = sorted(k for k in sections if re.match(rf"^{n}\.\d", k))
        sub_order.extend(subs)

    blocks: list[tuple[str, bool]] = []
    for main in order_main:
        if main in sections:
            blocks.append((main, True))
            if main.startswith("3.") or main.startswith("4.") or main.startswith("5."):
                prefix = main.split()[0]  # "3." etc - wrong

        num = main.split(".")[0] if main[0].isdigit() else None
        if num and num.isdigit():
            for sk in sub_order:
                if sk.startswith(f"{num}."):
                    blocks.append((sk, True))
                    for line in sections.get(sk, []):
                        blocks.append((line, False))
        else:
            for line in sections.get(main, []):
                blocks.append((line, False))

    # Fix: rebuild properly
    blocks = []
    for main in order_main:
        if main not in sections and main not in ("ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ"):
            continue
        if main in sections:
            blocks.append((main, True))
            num = None
            m = re.match(r"^(\d+)\.", main)
            if m:
                num = m.group(1)
            if num:
                for sk in sorted(k for k in sections if re.match(rf"^{num}\.\d", k)):
                    blocks.append((sk, True))
                    for line in sections[sk]:
                        blocks.append((line, False))
            else:
                for line in sections[main]:
                    blocks.append((line, False))
    return blocks


def main() -> None:
    sections = _parse_md(MD)
    doc = Document(str(SRC))

    # Introduction before chapter 1 (second occurrence — after TOC)
    ch1 = _find(doc, "1. Предпроектные исследования")
    intro_blocks = _ordered_blocks({"ВВЕДЕНИЕ": sections.get("ВВЕДЕНИЕ", [])})
    # only intro part
    intro_only: list[tuple[str, bool]] = [("ВВЕДЕНИЕ", True)]
    for line in sections.get("ВВЕДЕНИЕ", []):
        intro_only.append((line, False))

    if ch1 > 0:
        for text, is_h in reversed(intro_only):
            _insert_before(doc, ch1, text, bold=is_h)

    doc.save(str(OUT))
    doc = Document(str(OUT))

    # Chapters 3-5 + conclusion before literature list
    lit = _find(doc, "Список использованной литературы")
    ch3_blocks = _ordered_blocks(sections)
    # exclude intro from ch3_blocks
    ch3_blocks = [b for b in ch3_blocks if b[0] != "ВВЕДЕНИЕ"]

    if lit > 0:
        for text, is_h in reversed(ch3_blocks):
            _insert_before(doc, lit, text, bold=is_h)

    doc.save(str(OUT))
    print(f"OK: {OUT}")


if __name__ == "__main__":
    main()
