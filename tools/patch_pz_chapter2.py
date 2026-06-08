# -*- coding: utf-8 -*-
"""Restructure chapter 2 in ПЗ.docx."""
from __future__ import annotations

import copy
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "ПЗ" / "ПЗ.docx"
OUT = ROOT / "ПЗ" / "ПЗ.docx"

CH2_TOC = [
    "2.1 Язык программирования и архитектура системы",
    "2.2 Сбор и нормализация журналов Windows",
    "2.3 Хранение данных и правила детектирования",
    "2.4 Интерфейс, упаковка и эксплуатация",
    "2.5 Сравнение с альтернативами и итоговый стек технологий",
]

REMOVE_PATTERNS = [
    r"Дипломный проект выполнен студентом",
    r"Сведения об исполнителе",
    r"рекомендуется к защите",
]

CHAPTER_2_BLOCKS: list[tuple[str, bool]] = [
    (
        "Выбор технологий для дипломного mini SIEM определяется учебными целями "
        "(прозрачность кода, воспроизводимость стенда на одном ноутбуке), ограничениями "
        "доступа к корпоративной инфраструктуре оператора на время практики и инженерными "
        "компромиссами между простотой эксплуатации и достаточной выразительностью "
        "средств анализа событий.",
        False,
    ),
    ("2.1 Язык программирования и архитектура системы", True),
    (
        "В качестве основного языка выбран Python версии 3.11 (разрядность x64), "
        "рекомендованный в документации проекта для совместимости с используемыми "
        "колёсами и PyInstaller. Python обеспечивает высокую скорость разработки "
        "прототипов, богатую стандартную библиотеку для работы с текстом, временем, "
        "структурами данных и многопоточностью, а также зрелую интеграцию с Windows "
        "через пакет pywin32 для доступа к Win32 API журналов событий. Для задач "
        "сбора логов доминирует режим ожидания ввода-вывода; интерпретируемый язык "
        "при пакетном чтении и аккуратной записи в SQLite не становится узким местом.",
        False,
    ),
    (
        "Для mini SIEM выбран модульный стиль с разделением этапов «сбор → хранение → "
        "анализ → представление». Модульность отражает промышленную архитектуру SIEM, "
        "где сборщики и хранилища часто независимы и масштабируются различным образом. "
        "Декомпозиция выражается в отдельных модулях базы данных (db.py), правил "
        "(rules.py и YAML-файл), чтения журналов Windows (collector_win.py), импорта "
        "архивов (evtx_import.py), отчётности (report.py) и графического приложения "
        "(app.py). В операторской среде критичные логи находятся в зоне ограниченного "
        "доступа; для дипломного проекта целесообразен локальный контур, воспроизводимый "
        "на стенде без доступа к production.",
        False,
    ),
    ("2.2 Сбор и нормализация журналов Windows", True),
    (
        "Классические журналы Security, System и Application читаются через Win32-механизм; "
        "для канала Microsoft-Windows-Sysmon/Operational требуется API подсистемы Windows "
        "Event Log (Evt), поскольку «длинные» каналы не обслуживаются legacy-функциями "
        "OpenEventLog универсально. Нормализация идентификатора события с маской младших "
        "16 бит необходима для корректной работы правил, опирающихся на канонические "
        "номера событий Microsoft.",
        False,
    ),
    (
        "Качество детектирования зависит от стабильной и единообразной интерпретации "
        "полей событий при разных версиях Windows и разных источниках (классический "
        "журнал и Sysmon). В mini SIEM применён подход raw-normalized: события "
        "приводятся к общему набору полей (время, канал, provider, EventID, level, "
        "computer, record_id, message) и дополняются JSON-полем с исходными деталями. "
        "Подход сохраняет доказательную базу и обеспечивает универсальные SQL-запросы "
        "для фильтрации и правил.",
        False,
    ),
    (
        "Для событий входа Security 4624/4625 в аналитике значимы учётная запись, тип "
        "входа (в том числе удалённые сценарии) и IP-адрес источника. В Windows часть "
        "данных находится в массиве строковых вставок с форматом, зависящим от локали "
        "и версии. В проекте предусмотрено хранение сырого массива вставок (JSON) и "
        "отдельных разобранных полей (parsed_account, parsed_ip, parsed_logon_type) "
        "для ускорения запросов. Sysmon рассматривается как опциональный, но "
        "рекомендуемый источник для детектов PowerShell с EncodedCommand.",
        False,
    ),
    ("2.3 Хранение данных и правила детектирования", True),
    (
        "В качестве встраиваемого хранилища выбран SQLite: отсутствует необходимость "
        "отдельного сервера СУБД на учебном стенде; обеспечивается переносимость одного "
        "файла базы; SQL достаточен для агрегаций и оконных выборок. Режим журнала WAL "
        "снижает блокировки при одновременной записи сборщиком и чтении интерфейсом. "
        "Файл базы размещается в каталоге данных пользователя (LOCALAPPDATA\\miniSIEM). "
        "Отдельные таблицы events, alerts и actions отражают практику эксплуатации: "
        "фиксируются факт алерта, статусы обработки и след реакции аналитика.",
        False,
    ),
    (
        "Формат YAML выбран для человекочитаемости правил и версионирования в системе "
        "контроля версий. Правило содержит идентификатор, заголовок, описание, уровень "
        "серьёзности, имя детектора и блок параметров. Пороги чувствительности вынесены "
        "в параметры threshold и window_minutes. Для демонстрации на защите предусмотрены "
        "синтетические сценарии и импорт EVTX с заранее подготовленным набором событий. "
        "Для локального SQLite предусмотрены операции обслуживания: удаление записей "
        "старше заданного срока, checkpoint WAL и vacuum.",
        False,
    ),
    ("2.4 Интерфейс, упаковка и эксплуатация", True),
    (
        "Для десктопной Windows-среды выбран фреймворк Qt через биндинги PySide6: "
        "нативный внешний вид, табличные представления для больших объёмов событий, "
        "фильтрация по времени и каналу, поиск по EventID, просмотр алертов и смена "
        "статуса обработки. Веб-интерфейс увеличил бы объём фронтенд-разработки и "
        "усложнил демонстрацию без поднятого сервера.",
        False,
    ),
    (
        "PyInstaller позволяет собрать исполняемый файл miniSIEM.exe для переноса на "
        "машину без установленного Python. Стек разработки включает pywin32, PyYAML, "
        "sqlite3, PySide6; виртуальное окружение Python, requirements.txt и Git для "
        "воспроизводимости. Сбор событий и тяжёлые запросы к БД выполняются вне "
        "главного потока UI в отдельных потоках Qt.",
        False,
    ),
    (
        "Чтение журнала Security требует привилегий; реакции на алерты по умолчанию "
        "выполняются в режиме симуляции. Файл SQLite может содержать учётные имена и "
        "IP-адреса; доступ ограничивается политиками ОС. Демонстрации проводятся на "
        "тестовом стенде без данных реальных клиентов.",
        False,
    ),
    ("2.5 Сравнение с альтернативами и итоговый стек технологий", True),
    (
        "Elasticsearch или OpenSearch обеспечили бы полнотекстовый поиск и "
        "горизонтальное масштабирование, но потребовали бы JVM-кластера и отдельной "
        "эксплуатации — за рамками простоты дипломного стенда. PostgreSQL с "
        "партиционированием — компромисс для одного сервера; для single-user "
        "приложения SQLite остаётся прямолинейным выбором. Дипломный проект "
        "намеренно использует минимально достаточную инфраструктуру для снижения "
        "трудозатрат и рисков демонстрации.",
        False,
    ),
    (
        "PySide6 распространяется под LGPL; SQLite — public domain; Python — PSF "
        "license. Стек не требует закупки коммерческих лицензий SIEM. Перспективы "
        "развития: экспорт в syslog/JSON Lines, PostgreSQL, брокер сообщений, "
        "внешние списки IOC, доставка алертов по e-mail или Telegram.",
        False,
    ),
    ("Таблица 5 — Стек технологий mini SIEM (итоговая)", False),
    (
        "Версии зависимостей фиксируются в requirements.txt; используется виртуальное "
        "окружение для исключения дрейфа библиотек. Сборка выполняется под Windows 10/11 "
        "x64; pywin32 и PyInstaller зависят от платформы. Сообщения журналов Windows "
        "могут быть локализованы; детекторы по подстрокам командной строки требуют "
        "учёта языка интерфейса на учебном стенде.",
        False,
    ),
]


def _p_text(p) -> str:
    return "".join(r.text for r in p.runs).strip()


def _style_para(p, *, bold: bool = False) -> None:
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
        if bold:
            r.bold = True


def _delete_paragraph(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def _clear_between(start_paragraph, end_paragraph) -> None:
    body = start_paragraph._element.getparent()
    start = start_paragraph._element
    end = end_paragraph._element
    removing = False
    to_remove = []
    for child in list(body):
        if child is start:
            removing = True
            continue
        if child is end:
            break
        if removing:
            to_remove.append(child)
    for child in to_remove:
        body.remove(child)


def _find_stack_table_element(start_paragraph, end_paragraph):
    start = start_paragraph._element
    end = end_paragraph._element
    body = start.getparent()
    removing = False
    for child in body:
        if child is start:
            removing = True
            continue
        if child is end:
            break
        if removing and child.tag.endswith("tbl"):
            all_text = "".join(
                (node.text or "") for node in child.iter() if node.tag.endswith("}t")
            )
            if "Компонент стека" in all_text:
                return copy.deepcopy(child)
    return None


def _insert_element_before(anchor_paragraph, element) -> None:
    parent = anchor_paragraph._element.getparent()
    parent.insert(parent.index(anchor_paragraph._element), element)


def _should_remove(text: str) -> bool:
    return any(re.search(pat, text, re.I) for pat in REMOVE_PATTERNS)


def _find_body_intro_idx(doc: Document) -> int:
    for i, p in enumerate(doc.paragraphs):
        t = _p_text(p)
        if t == "Введение":
            for j in range(i + 1, min(i + 4, len(doc.paragraphs))):
                nxt = _p_text(doc.paragraphs[j])
                if len(nxt) > 50:
                    return i
    return 0


def _find_ch2_idx(doc: Document, after: int) -> int:
    for i in range(after + 1, len(doc.paragraphs)):
        if re.match(r"^2[\.\s]+Выбор технологий", _p_text(doc.paragraphs[i])):
            return i
    raise RuntimeError("Chapter 2 not found")


def _find_ch3_idx(doc: Document, after: int) -> int:
    for i in range(after + 1, len(doc.paragraphs)):
        if re.match(r"^3[\.\s]+Проектирование", _p_text(doc.paragraphs[i])):
            return i
    raise RuntimeError("Chapter 3 not found")


def _insert_before(doc: Document, anchor_paragraph, text: str, *, bold: bool = False) -> None:
    anchor = anchor_paragraph._element
    parent = anchor.getparent()
    new_p = OxmlElement("w:p")
    parent.insert(parent.index(anchor), new_p)
    para = Paragraph(new_p, doc._body)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = bold
    _style_para(para, bold=bold)


def main() -> None:
    src = SRC if SRC.exists() else ROOT / "ПЗ" / "ПЗ.docx"
    doc = Document(str(src))

    for i in range(len(doc.paragraphs) - 1, -1, -1):
        t = _p_text(doc.paragraphs[i])
        if t and _should_remove(t):
            _delete_paragraph(doc.paragraphs[i])

    intro = _find_body_intro_idx(doc)
    ch2 = _find_ch2_idx(doc, intro)
    ch3 = _find_ch3_idx(doc, ch2)

    ch3 = _find_ch3_idx(doc, ch2)
    p2 = doc.paragraphs[ch2]
    p3 = doc.paragraphs[ch3]
    stack_table = _find_stack_table_element(p2, p3)

    for r in p2.runs:
        r.text = ""
    if p2.runs:
        p2.runs[0].text = "2 Выбор технологий и средств разработки"
    else:
        p2.add_run("2 Выбор технологий и средств разработки")
    _style_para(p2, bold=True)

    _clear_between(p2, p3)

    ch3 = _find_ch3_idx(doc, ch2)
    anchor = doc.paragraphs[ch3]
    table_inserted = False
    for text, bold in CHAPTER_2_BLOCKS:
        _insert_before(doc, anchor, text, bold=bold)
        if stack_table is not None and text.startswith("Таблица 5"):
            _insert_element_before(anchor, stack_table)
            table_inserted = True
    if stack_table is not None and not table_inserted:
        _insert_element_before(anchor, stack_table)

    # TOC (если есть блок содержания в начале документа)
    for i, p in enumerate(doc.paragraphs[:40]):
        t = _p_text(p)
        if re.match(r"^2[\.\s]*Выбор технологий", t) and i < 25:
            nxt = _p_text(doc.paragraphs[i + 1]) if i + 1 < len(doc.paragraphs) else ""
            if not nxt.startswith("2.1"):
                ap = doc.paragraphs[i + 1]
                for sub in reversed(CH2_TOC):
                    _insert_before(doc, ap, sub, bold=False)
            for r in doc.paragraphs[i].runs:
                r.text = ""
            if doc.paragraphs[i].runs:
                doc.paragraphs[i].runs[0].text = "2 Выбор технологий и средств разработки"
            break

    out_path = OUT
    try:
        doc.save(str(out_path))
    except PermissionError:
        out_path = ROOT / "ПЗ" / "ПЗ_обновлено.docx"
        doc.save(str(out_path))
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
