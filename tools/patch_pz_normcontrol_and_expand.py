# -*- coding: utf-8 -*-
"""Нормконтроль (таблицы/рисунки 1.1, 3.1), места для картинок, +~7 стр. текста."""
from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "tools"))
from pz_expansion_extra3 import EXPANSIONS_EXTRA3  # noqa: E402
from pz_expansion_extra5 import EXPANSIONS_EXTRA5  # noqa: E402
from pz_expansion_extra6 import EXPANSIONS_EXTRA6  # noqa: E402

SRC = ROOT / "ПЗ" / "ПЗ.docx"
OUT = ROOT / "ПЗ" / "ПЗ.docx"
BACKUP = ROOT / "ПЗ" / "ПЗ_before_normcontrol.docx"

CAPTION_RENAMES: dict[str, str] = {
    "Таблица 1 —": "Таблица 1.1 –",
    "Таблица 2 —": "Таблица 1.2 –",
    "Таблица 3 —": "Таблица 1.3 –",
    "Таблица 4 —": "Таблица 1.4 –",
    "Таблица 5 —": "Таблица 2.1 –",
    "Таблица 6 —": "Таблица 3.1 –",
    "Таблица 7 —": "Таблица 3.2 –",
    "Таблица 8 —": "Таблица 4.1 –",
    "Таблица 9 —": "Таблица 5.1 –",
    "Рисунок 1 —": "Рисунок 3.1 –",
}

TEXT_RENAMES: list[tuple[str, str]] = [
    ("Таблица 9 фиксирует", "Таблица 5.1 фиксирует"),
    ("Таблица 7 перечисляет", "Таблица 3.2 перечисляет"),
    ("согласно таблице 9", "согласно таблице 5.1"),
    ("приведена в таблице 4", "приведена в таблице 1.4"),
    ("(рисунок 1)", "(рисунок 3.1)"),
    ("рисунке 1", "рисунке 3.1"),
    ("«Рисунок N —", "«Рисунок N –"),
    ("таблице 9", "таблице 5.1"),
    ("таблице 7", "таблице 3.2"),
    ("таблице 6", "таблице 3.1"),
    ("таблице 5", "таблице 2.1"),
    ("таблице 4", "таблице 1.4"),
    ("таблице 3", "таблице 1.3"),
    ("таблице 2", "таблице 1.2"),
    ("таблице 1", "таблице 1.1"),
]

ANCHORS: list[tuple[str, str]] = [
    ("Введение", "1 Предпроектные"),
    ("1.1", "1.2"),
    ("1.2", "1.3"),
    ("1.3", "2 Выбор"),
    ("2.1", "2.2"),
    ("2.2", "2.3"),
    ("2.3", "2.4"),
    ("2.4", "2.5"),
    ("2.5", "3 Проектирование"),
    ("3.1", "3.2"),
    ("3.2", "3.3"),
    ("3.3", "3.4"),
    ("3.4", "3.5"),
    ("3.5", "3.6"),
    ("3.6", "4 Реализация"),
    ("4.1", "4.2"),
    ("4.2", "4.3"),
    ("4.3", "4.4"),
    ("4.4", "4.5"),
    ("4.5", "5 Тестирование"),
    ("5.1", "5.2"),
    ("5.2", "5.3"),
    ("5.3", "5.4"),
    ("5.4", "5.5"),
    ("5.5", "Заключение"),
    ("Заключение", "Список использованной"),
]

EXTRA_BLOCKS: dict[str, list[str]] = {}
for _src in (EXPANSIONS_EXTRA3, EXPANSIONS_EXTRA5, EXPANSIONS_EXTRA6):
    for k, paras in _src.items():
        EXTRA_BLOCKS.setdefault(k, []).extend(paras)

FIGURES: list[dict] = [
    {
        "after_snippet": "ER-диаграмма (графический материал)",
        "ref": "Структура связей между сущностями базы данных приведена на рисунке 3.2.",
        "caption": "Рисунок 3.2 – ER-диаграмма базы данных mini SIEM",
    },
    {
        "after_snippet": "Подсистема импорта EVTX",
        "ref": "Последовательность вызовов при поступлении новой порции событий показана на рисунке 3.3.",
        "caption": "Рисунок 3.3 – Диаграмма последовательности обработки события",
    },
    {
        "after_snippet": "Проектирование UX учитывает типовой сценарий",
        "ref": "Сводная панель Dashboard программного комплекса представлена на рисунке 3.4.",
        "caption": "Рисунок 3.4 – Вкладка Dashboard программного комплекса mini SIEM",
    },
    {
        "after_snippet": "Класс _EventsTab реализует панель фильтров",
        "ref": "Интерфейс просмотра и фильтрации событий показан на рисунке 4.1.",
        "caption": "Рисунок 4.1 – Вкладка Events с панелью фильтров",
    },
    {
        "after_snippet": "_AlertsTab поддерживает смену статуса",
        "ref": "Список алертов с уровнями серьёзности и статусами обработки приведён на рисунке 4.2.",
        "caption": "Рисунок 4.2 – Вкладка Alerts программного комплекса mini SIEM",
    },
    {
        "after_snippet": "Функция load_rule_defs",
        "ref": "Управление правилами детектирования и параметрами порогов показано на рисунке 4.3.",
        "caption": "Рисунок 4.3 – Вкладка Rules с параметрами правил",
    },
    {
        "after_snippet": "`report.py` генерирует HTML",
        "ref": "Пример HTML-отчёта по инциденту представлен на рисунке 4.4.",
        "caption": "Рисунок 4.4 – HTML-отчёт по инциденту",
    },
    {
        "after_snippet": "Импорт реализован в `_EvtxImportWorker`",
        "ref": "Диалог импорта архива EVTX с индикацией прогресса показан на рисунке 4.5.",
        "caption": "Рисунок 4.5 – Импорт файла EVTX",
    },
    {
        "after_snippet": "5.3 Тестирование правил детектирования",
        "ref": "Пример срабатывания правила детектирования на тестовых данных приведён на рисунке 5.1.",
        "caption": "Рисунок 5.1 – Срабатывание правила WIN-SEC-4625-BURST",
    },
]

ASCII_MARKERS = (
    "[ Windows Event Log",
    "[ collector_win",
    "[ Нормализация ]",
    "[ db.py / SQLite ]",
    "[ rules.py ]",
    "[ alerts ]",
    "[ actions / dry-run ]",
    "+-----+-----+",
    "|           |",
    "v           v",
    "-->",
)


def _p_text(p) -> str:
    return "".join(r.text for r in p.runs).strip()


def _norm(s: str) -> str:
    s = s.lower().replace("ё", "е")
    return re.sub(r"\s+", " ", s).strip()


def _style_body(p, *, center: bool = False, indent: bool = True) -> None:
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(0) if center else (Cm(1.25) if indent else Cm(0))
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)


def _insert_after(doc: Document, anchor_paragraph, text: str, *, center: bool = False) -> Paragraph:
    anchor = anchor_paragraph._element
    parent = anchor.getparent()
    new_p = OxmlElement("w:p")
    parent.insert(parent.index(anchor) + 1, new_p)
    para = Paragraph(new_p, doc._body)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    _style_body(para, center=center, indent=not center)
    return para


def _insert_before(doc: Document, anchor_paragraph, text: str, *, center: bool = False) -> Paragraph:
    anchor = anchor_paragraph._element
    parent = anchor.getparent()
    new_p = OxmlElement("w:p")
    parent.insert(parent.index(anchor), new_p)
    para = Paragraph(new_p, doc._body)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    _style_body(para, center=center, indent=not center)
    return para


def _replace_para_text(p, new_text: str) -> None:
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(new_text)


def _apply_text_renames(text: str) -> str:
    for old, new in TEXT_RENAMES:
        text = text.replace(old, new)
    for old, new in CAPTION_RENAMES.items():
        if text.startswith(old):
            text = new + text[len(old) :]
    return text


def _is_caption_line(t: str) -> bool:
    return bool(re.match(r"^(Таблица \d+\.\d|Таблица \d+ —|Рисунок \d)", t))


def _rename_captions(doc: Document) -> int:
    n = 0
    for p in doc.paragraphs:
        t = _p_text(p)
        if not t:
            continue
        for old, new in CAPTION_RENAMES.items():
            if t.startswith(old):
                _replace_para_text(p, new + t[len(old) :])
                _style_body(p, center=True, indent=False)
                n += 1
                break
    return n


def _rename_body_refs(doc: Document) -> int:
    n = 0
    for p in doc.paragraphs:
        t = _p_text(p)
        if not t:
            continue
        if re.match(r"^(Таблица \d+\.\d –|Рисунок \d+\.\d –)", t):
            continue
        new_t = _apply_text_renames(t)
        if new_t != t:
            _replace_para_text(p, new_t)
            n += 1
    return n


def _remove_ascii_diagram(doc: Document) -> int:
    to_remove = []
    for p in doc.paragraphs:
        t = _p_text(p)
        if not t:
            continue
        if any(t.startswith(m) or m in t for m in ASCII_MARKERS):
            to_remove.append(p)
        elif t in ("|", "v", "^"):
            to_remove.append(p)
    for p in to_remove:
        p._element.getparent().remove(p._element)
    return len(to_remove)


def _find_para(doc: Document, snippet: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if snippet in _p_text(p):
            return p
    return None


def _find_anchor(doc: Document, prefix: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if _p_text(p).startswith(prefix):
            return p
    return None


def _fix_figure_31(doc: Document) -> None:
    existing = {_norm(_p_text(p)) for p in doc.paragraphs}
    for p in doc.paragraphs:
        t = _p_text(p)
        if not (t.startswith("Рисунок 1 —") or t.startswith("Рисунок 3.1")):
            continue
        ref = "логическая архитектура комплекса mini siem представлена на рисунке 3.1"
        if ref not in existing:
            _insert_before(
                doc,
                p,
                "Логическая архитектура комплекса mini SIEM представлена на рисунке 3.1.",
            )
            _insert_before(doc, p, "[Место для рисунка]", center=True)
        _replace_para_text(p, "Рисунок 3.1 – Архитектура mini SIEM (логические компоненты)")
        _style_body(p, center=True, indent=False)
        break
    _remove_ascii_diagram(doc)


def _add_figure_block(doc: Document, anchor_paragraph, ref: str, caption: str) -> None:
    """Порядок по нормконтролю: текст-ссылка → место → подпись."""
    p = _insert_after(doc, anchor_paragraph, ref)
    p = _insert_after(doc, p, "[Место для рисунка]", center=True)
    _insert_after(doc, p, caption, center=True)


def _add_figure_placeholders(doc: Document) -> int:
    added = 0
    existing_caps = {_norm(_p_text(p)) for p in doc.paragraphs if _p_text(p)}
    for fig in FIGURES:
        anchor = _find_para(doc, fig["after_snippet"])
        if anchor is None:
            print(f"WARN figure: {fig['after_snippet'][:55]}")
            continue
        cap = _norm(fig["caption"])
        if cap in existing_caps:
            continue
        _add_figure_block(doc, anchor, fig["ref"], fig["caption"])
        existing_caps.add(cap)
        added += 1
    return added


def _add_expansion_text(doc: Document, existing: set[str], target_add: int = 20_000) -> int:
    inserted = 0
    added_chars = 0
    skip_starts = ("Глава ", "Раздел ", "Таблица 1 сравнивает", "Логическая схема на рисунке 1")
    for key, next_prefix in ANCHORS:
        if added_chars >= target_add:
            break
        anchor = _find_anchor(doc, next_prefix)
        if anchor is None:
            continue
        count = 0
        for text in EXTRA_BLOCKS.get(key, []):
            if added_chars >= target_add or count >= 2:
                break
            if any(text.startswith(s) for s in skip_starts):
                continue
            text = _apply_text_renames(text)
            n = _norm(text)
            if n in existing or len(text) < 80:
                continue
            _insert_before(doc, anchor, text)
            existing.add(n)
            added_chars += len(text)
            inserted += 1
            count += 1
    return inserted


def main() -> None:
    shutil.copy2(SRC, BACKUP)
    doc = Document(str(SRC))
    before = sum(len(_p_text(p)) for p in doc.paragraphs if _p_text(p))

    caps = _rename_captions(doc)
    refs = _rename_body_refs(doc)
    _fix_figure_31(doc)
    figs = _add_figure_placeholders(doc)

    existing = {_norm(_p_text(p)) for p in doc.paragraphs if _p_text(p)}
    added = _add_expansion_text(doc, existing, target_add=20_000)

    try:
        doc.save(str(OUT))
    except PermissionError:
        alt = ROOT / "ПЗ" / "ПЗ_normcontrol.docx"
        doc.save(str(alt))
        print(f"WARN: {SRC.name} locked, saved {alt.name}")
        OUT_PATH = alt
    else:
        OUT_PATH = OUT

    after = sum(len(_p_text(p)) for p in doc.paragraphs if _p_text(p))
    print(f"Captions renamed: {caps}, body refs: {refs}, figures: {figs}, +paras: {added}")
    print(f"Chars: {before} -> {after} (+{after - before})")
    print(f"Backup: {BACKUP.name}")
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
