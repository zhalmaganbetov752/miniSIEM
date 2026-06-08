# -*- coding: utf-8 -*-
"""Исправление порядка рисунков и донабор ~7 стр. текста."""
from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "tools"))
from pz_expansion_extra2 import EXPANSIONS_EXTRA2  # noqa: E402
from pz_expansion_extra4 import EXPANSIONS_EXTRA4  # noqa: E402

SRC = ROOT / "ПЗ" / "ПЗ.docx"
OUT = ROOT / "ПЗ" / "ПЗ.docx"

ANCHORS: list[tuple[str, str]] = [
    ("1.1", "1.2"), ("1.2", "1.3"), ("1.3", "2 Выбор"),
    ("2.1", "2.2"), ("2.2", "2.3"), ("2.3", "2.4"), ("2.4", "2.5"),
    ("3.1", "3.2"), ("3.2", "3.3"), ("3.3", "3.4"), ("3.4", "3.5"), ("3.5", "3.6"),
    ("4.1", "4.2"), ("4.2", "4.3"), ("4.3", "4.4"), ("4.4", "4.5"),
    ("5.1", "5.2"), ("5.2", "5.3"), ("5.3", "5.4"), ("5.4", "5.5"), ("5.5", "Заключение"),
]

SKIP = ("Глава ", "Раздел ", "Таблица 1 сравнивает", "Логическая схема на рисунке 1", "Рисунки 2–7")


def _p_text(p) -> str:
    return "".join(r.text for r in p.runs).strip()


def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", s.lower().replace("ё", "е")).strip()


def _style_body(p, *, center: bool = False) -> None:
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(0) if center else Cm(1.25)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)


def _insert_after(doc: Document, anchor: Paragraph, text: str, *, center: bool = False) -> Paragraph:
    el = anchor._element
    parent = el.getparent()
    new_p = OxmlElement("w:p")
    parent.insert(parent.index(el) + 1, new_p)
    para = Paragraph(new_p, doc._body)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    _style_body(para, center=center)
    return para


def _remove_para(p: Paragraph) -> None:
    p._element.getparent().remove(p._element)


def _repair_figure_triplets(doc: Document) -> int:
    """caption → placeholder → ref  =>  ref → placeholder → caption"""
    fixed = 0
    paras = list(doc.paragraphs)
    i = 0
    while i < len(paras) - 2:
        t0, t1, t2 = _p_text(paras[i]), _p_text(paras[i + 1]), _p_text(paras[i + 2])
        if (
            re.match(r"^Рисунок \d+\.\d+ –", t0)
            and t1 == "[Место для рисунка]"
            and "рисунке" in t2.lower()
        ):
            anchor = paras[i - 1] if i > 0 else paras[i]
            ref, cap = t2, t0
            for p in (paras[i], paras[i + 1], paras[i + 2]):
                _remove_para(p)
            p = _insert_after(doc, anchor, ref)
            p = _insert_after(doc, p, "[Место для рисунка]", center=True)
            _insert_after(doc, p, cap, center=True)
            fixed += 1
            paras = list(doc.paragraphs)
            i = 0
            continue
        i += 1
    return fixed


def _fix_gost_caption_text(doc: Document) -> None:
    for p in doc.paragraphs:
        t = _p_text(p)
        if "«Рисунок N —" in t or "«Рисунок N –" in t:
            _replace(p, t.replace("«Рисунок N —", "«Рисунок X.Y –").replace("«Рисунок N –", "«Рисунок X.Y –"))
        if "оформляются по ГОСТ" in t and "—" in t:
            _replace(p, t.replace("«Рисунок N — Название»", "«Рисунок 3.1 – Название»"))


def _replace(p: Paragraph, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def _find_anchor(doc: Document, prefix: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if _p_text(p).startswith(prefix):
            return p
    return None


def _insert_before(doc: Document, anchor: Paragraph, text: str) -> None:
    el = anchor._element
    parent = el.getparent()
    new_p = OxmlElement("w:p")
    parent.insert(parent.index(el), new_p)
    para = Paragraph(new_p, doc._body)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    _style_body(para)


def _add_more_text(doc: Document, target: int = 8000) -> int:
    blocks: dict[str, list[str]] = {}
    for src in (EXPANSIONS_EXTRA2, EXPANSIONS_EXTRA4):
        for k, v in src.items():
            blocks.setdefault(k, []).extend(v)
    existing = {_norm(_p_text(p)) for p in doc.paragraphs if _p_text(p)}
    added = 0
    chars = 0
    for key, nxt in ANCHORS:
        if chars >= target:
            break
        anchor = _find_anchor(doc, nxt)
        if not anchor:
            continue
        n = 0
        for text in blocks.get(key, []):
            if chars >= target or n >= 2:
                break
            if any(text.startswith(s) for s in SKIP):
                continue
            text = text.replace("таблице 4", "таблице 1.4").replace("Таблица 6 ", "Таблица 3.1 ")
            text = text.replace("рисунке 1", "рисунке 3.1").replace("Таблица 9", "Таблица 5.1")
            nn = _norm(text)
            if nn in existing or len(text) < 80:
                continue
            _insert_before(doc, anchor, text)
            existing.add(nn)
            chars += len(text)
            added += 1
            n += 1
    return added


def main() -> None:
    doc = Document(str(SRC))
    before = sum(len(_p_text(p)) for p in doc.paragraphs if _p_text(p))
    fixed = _repair_figure_triplets(doc)
    _fix_gost_caption_text(doc)
    more = _add_more_text(doc, target=8000)
    doc.save(str(OUT))
    after = sum(len(_p_text(p)) for p in doc.paragraphs if _p_text(p))
    print(f"Fixed figure blocks: {fixed}, +paragraphs: {more}")
    print(f"Chars: {before} -> {after} (+{after - before})")


if __name__ == "__main__":
    main()
